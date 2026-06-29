"""
Build the client-facing intake checklist DOCX for RequiredFilings.com.
Self-explanatory document for the client to fill in.
Generates: client-intake-checklist.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----- colour and style helpers -----
NAVY = RGBColor(0x0A, 0x19, 0x29)
ACCENT = RGBColor(0x1F, 0x4E, 0x9C)
GREY = RGBColor(0x55, 0x5B, 0x6A)
LIGHT_GREY = RGBColor(0xC9, 0xCC, 0xD3)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, color="C9CCD3", size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_field_row(table, label, hint=""):
    """A two-row block: label + hint, then a blank line for the answer."""
    row = table.add_row()
    cell = row.cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    if hint:
        p2 = cell.add_paragraph()
        rh = p2.add_run(hint)
        rh.italic = True
        rh.font.size = Pt(9)
        rh.font.color.rgb = GREY
    set_cell_bg(cell, "F4F5F7")
    set_cell_border(cell)

    answer_row = table.add_row()
    a = answer_row.cells[0]
    a.text = ""
    ap = a.paragraphs[0]
    ap.paragraph_format.space_before = Pt(6)
    ap.paragraph_format.space_after = Pt(6)
    ap.add_run("                                                                                                                                          ")
    set_cell_border(a)


def add_field_table(doc):
    table = doc.add_table(rows=0, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16)
    return table


def add_heading_block(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"SECTION {number}   ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = ACCENT
    r2 = p.add_run(title.upper())
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = NAVY
    # underline rule
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    rr = rule.add_run("_" * 95)
    rr.font.color.rgb = LIGHT_GREY
    rr.font.size = Pt(8)


def add_intro_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def add_checkbox_line(doc, label, hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    box = p.add_run("[   ]   ")
    box.font.size = Pt(11)
    box.bold = True
    txt = p.add_run(label)
    txt.font.size = Pt(10)
    txt.font.color.rgb = NAVY
    if hint:
        h = p.add_run(f"   ({hint})")
        h.italic = True
        h.font.size = Pt(9)
        h.font.color.rgb = GREY


# ===== build the document =====
doc = Document()

# global font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# page margins
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


# ---------- COVER ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(2)
rt = title.add_run("WEBSITE LAUNCH INTAKE")
rt.bold = True
rt.font.size = Pt(11)
rt.font.color.rgb = ACCENT

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(2)
rs = subtitle.add_run("RequiredFilings.com")
rs.bold = True
rs.font.size = Pt(28)
rs.font.color.rgb = NAVY

tag = doc.add_paragraph()
tag.paragraph_format.space_after = Pt(18)
rtag = tag.add_run("Information required to take the website live")
rtag.italic = True
rtag.font.size = Pt(12)
rtag.font.color.rgb = GREY

# meta box
meta = doc.add_table(rows=1, cols=2)
meta.autofit = False
meta.columns[0].width = Cm(8)
meta.columns[1].width = Cm(8)
meta_cells = meta.rows[0].cells

m1 = meta_cells[0]
m1.text = ""
mp = m1.paragraphs[0]
r = mp.add_run("PREPARED FOR\n")
r.bold = True; r.font.size = Pt(8); r.font.color.rgb = GREY
r2 = mp.add_run("Client / Business Owner\n")
r2.font.size = Pt(10); r2.font.color.rgb = NAVY
r3 = mp.add_run("RequiredFilings.com")
r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = NAVY
set_cell_bg(m1, "F4F5F7")
set_cell_border(m1)

m2 = meta_cells[1]
m2.text = ""
mp2 = m2.paragraphs[0]
r = mp2.add_run("DATE OF VISIT\n")
r.bold = True; r.font.size = Pt(8); r.font.color.rgb = GREY
r2 = mp2.add_run("_______ / _______ / 20_______\n\n")
r2.font.size = Pt(11); r2.font.color.rgb = NAVY
r3 = mp2.add_run("FILLED BY (NAME)\n")
r3.bold = True; r3.font.size = Pt(8); r3.font.color.rgb = GREY
r4 = mp2.add_run("_________________________________")
r4.font.size = Pt(11); r4.font.color.rgb = NAVY
set_cell_bg(m2, "F4F5F7")
set_cell_border(m2)


# ---------- ABOUT THIS DOCUMENT ----------
add_heading_block(doc, "A", "About this document")
add_intro_para(doc,
    "This is the final information request before your website goes live. "
    "Your website at RequiredFilings.com is built and ready. The pages, "
    "design, copy, legal text, and contact forms are all in place. The "
    "only thing left is to plug in your real business details and a few "
    "decisions that only you can make."
)
add_intro_para(doc,
    "How to fill this in: write your answers in the blank space below each "
    "question. If you do not have an answer yet, leave it blank and write "
    "\"pending\" next to it. For decisions where you see checkboxes [   ], "
    "tick the one you prefer. Sign the last page when complete. Our team "
    "will then update the website and schedule the go-live date."
)
add_intro_para(doc,
    "Estimated time to complete: 45 to 60 minutes if you have your "
    "registration certificates, GSTIN, and team details handy."
)


# ---------- WEBSITE OVERVIEW ----------
add_heading_block(doc, "B", "About your website")
add_intro_para(doc,
    "The website covers 15 pages: a homepage, about, contact, a services "
    "overview, 10 individual service pages (Start a Business, Statutory "
    "Registrations, Licenses, ROC Compliance, Tax Services, GST Services, "
    "IPR, ISO Certification, MSME ZED, Accounting), a blog index, and four "
    "legal pages (Privacy, Terms, Refund, Disclaimer). Roughly 32 percent "
    "of the launch work is complete. The remaining 68 percent is split "
    "between the information requested in this document and the technical "
    "go-live tasks our team will handle after you return this form."
)


# ===================================================================
# SECTION 1 — BUSINESS CONTACT DETAILS
# ===================================================================
add_heading_block(doc, "1", "Business contact details")
add_intro_para(doc,
    "These details appear in the website header, footer, contact page, "
    "every service page, and in your legal documents. They must be the "
    "real numbers and addresses your customers should reach you on."
)
t = add_field_table(doc)
add_field_row(t, "1.1   Primary business phone number",
              "Include country code. Example: +91 98765 43210. Used for click-to-call across the site.")
add_field_row(t, "1.2   WhatsApp number (for click-to-chat)",
              "Same as above or different? Include country code. Leave blank if WhatsApp is not offered.")
add_field_row(t, "1.3   Primary business email",
              "Example: hello@requiredfilings.com. Used on contact forms, footer, and legal pages.")
add_field_row(t, "1.4   Secondary email (optional)",
              "If you want a separate billing or support email. Example: accounts@requiredfilings.com.")
add_field_row(t, "1.5   Full office address (one line)",
              "Building, street, area, city, state, PIN code. Appears in header pop-out and footer.")
add_field_row(t, "1.6   Registered office address (if different)",
              "Used in your Privacy Policy and Terms of Service. Leave blank if same as office address.")
add_field_row(t, "1.7   City of registered office",
              "City name only. Used in the legal jurisdiction clause on Terms of Service.")
add_field_row(t, "1.8   Business working hours",
              "Example: Monday to Saturday, 10:00 AM to 7:00 PM IST. Shown on the contact page.")


# ===================================================================
# SECTION 2 — LEGAL & REGISTRATION DETAILS
# ===================================================================
add_heading_block(doc, "2", "Legal and registration details")
add_intro_para(doc,
    "These appear in the footer of every page and in the Terms and "
    "Disclaimer pages. Required for legal compliance and customer trust."
)
t = add_field_table(doc)
add_field_row(t, "2.1   Legal entity name",
              "The exact registered name of the company or firm that operates the business. "
              "Example: ABC Compliance Services Private Limited.")
add_field_row(t, "2.2   Business structure",
              "Tick one: Private Limited / LLP / OPC / Partnership / Proprietorship / Other (specify).")
add_field_row(t, "2.3   CIN (for companies) OR Firm Registration Number",
              "21-character CIN for Pvt Ltd, LLPIN for LLP, or registration number for other entities.")
add_field_row(t, "2.4   GSTIN",
              "15-character GST identification number. Appears in the footer of every page.")
add_field_row(t, "2.5   PAN of the entity",
              "10-character PAN. Used only for legal documents, not displayed on the site.")


# ===================================================================
# SECTION 3 — FOUNDER & TEAM
# ===================================================================
add_heading_block(doc, "3", "Founder and team")
add_intro_para(doc,
    "Used on the About page. Customers buy trust before service, so real "
    "names and faces matter. We need a short bio + photo for each person "
    "you want on the team section."
)
t = add_field_table(doc)
add_field_row(t, "3.1   Founder name",
              "Full name as you want it shown.")
add_field_row(t, "3.2   Year the business was founded",
              "Four-digit year. Example: 2018.")
add_field_row(t, "3.3   Short founder story (2 to 3 sentences)",
              "Why you started the business and who you built it for. Plain English, no marketing fluff.")

add_intro_para(doc, "Team members (up to 4 are displayed on the About page):")
for i in range(1, 5):
    t = add_field_table(doc)
    add_field_row(t, f"3.{3+i}   Team member {i} — full name")
    add_field_row(t, f"3.{3+i}a  Role / designation",
                  "Example: Senior Chartered Accountant, GST Specialist, Client Manager.")
    add_field_row(t, f"3.{3+i}b  LinkedIn profile URL (optional)")
    add_field_row(t, f"3.{3+i}c  Photo provided? (Y/N)",
                  "Please share a clear, front-facing portrait photo. Plain background preferred.")


# ===================================================================
# SECTION 4 — TRUST PROOF (TESTIMONIALS, STATS, CASE STUDIES)
# ===================================================================
add_heading_block(doc, "4", "Trust proof — testimonials, stats, case studies")
add_intro_para(doc,
    "Specificity builds trust. We need real numbers and real client names "
    "(with permission). Do not invent these. If you do not have data for "
    "a field, leave it blank and we will hide that section until you do."
)

# Stats
add_intro_para(doc, "Headline statistics (shown on homepage and about page):")
t = add_field_table(doc)
add_field_row(t, "4.1   Total filings completed to date",
              "Approximate number is fine. Example: 1,200+ or 5,000+.")
add_field_row(t, "4.2   On-time filing rate (percentage)",
              "Example: 98%. Be honest — customers will hold you to this.")
add_field_row(t, "4.3   Number of active clients",
              "Optional. Example: 250+ businesses served.")
add_field_row(t, "4.4   Years in operation",
              "Calculated from the year founded above. Confirm or override.")
add_field_row(t, "4.5   Government fee threshold for company registration",
              "The authorised share capital amount up to which government fees are included. "
              "Example: government fees included for authorised capital up to Rs. 1,00,000.")

# Testimonials
add_intro_para(doc, "")
add_intro_para(doc,
    "Testimonials (we display 3 on the homepage and up to 5 on the about page). "
    "Permission to use the client's name and business type is required. "
    "Fill in as many as you have:"
)
for i in range(1, 6):
    t = add_field_table(doc)
    add_field_row(t, f"4.6.{i}   Testimonial {i} — client's full name")
    add_field_row(t, f"4.6.{i}a  Business type and city",
                  "Example: Software startup founder, Bengaluru.")
    add_field_row(t, f"4.6.{i}b  Their quote (2 to 4 sentences in their own words)",
                  "Plain English. Specific outcome is better than \"great service\".")
    add_field_row(t, f"4.6.{i}c  Written permission to publish? (Y/N)",
                  "We need this in writing — email or signed note is fine.")

# Case studies
add_intro_para(doc, "")
add_intro_para(doc,
    "Case study outcomes (shown on the homepage as 3 short result cards):"
)
t = add_field_table(doc)
add_field_row(t, "4.7.1   SaaS / startup case — outcome line",
              "Currently reads: \"SaaS startup: Pvt Ltd registered and first-year ROC filed.\" "
              "Add the specific outcome metric. Example: completed in 14 days, ROC filed before due date.")
add_field_row(t, "4.7.2   Trader / small business case — outcome line",
              "Currently reads: \"Trader: two years of pending GST returns filed and notices cleared.\" "
              "Add the specific outcome. Example: GSTR-3B and GSTR-1 backfiled, penalty reduced by 60%.")
add_field_row(t, "4.7.3   Manufacturer case — outcome line",
              "Currently reads: \"Manufacturer: ISO 9001 certified and MSME Udyam approved.\" "
              "Add the specific outcome. Example: certified in 45 days, Udyam approved in 7 days.")


# ===================================================================
# SECTION 5 — SOCIAL MEDIA
# ===================================================================
add_heading_block(doc, "5", "Social media links")
add_intro_para(doc,
    "Appear as icons in the website footer. Leave blank for any platform "
    "you do not actively use — empty profiles hurt trust more than missing icons."
)
t = add_field_table(doc)
add_field_row(t, "5.1   LinkedIn company page URL",
              "Recommended — most important for a B2B compliance business.")
add_field_row(t, "5.2   Facebook page URL")
add_field_row(t, "5.3   Instagram profile URL")
add_field_row(t, "5.4   YouTube channel URL")
add_field_row(t, "5.5   Twitter / X profile URL")


# ===================================================================
# SECTION 6 — SERVICE PRICING
# ===================================================================
add_heading_block(doc, "6", "Service pricing")
add_intro_para(doc,
    "Each service page shows three pricing tiers. Fill in the rupee amount for "
    "each. If a service is quote-based, write \"On request\" and we will hide the "
    "price. All amounts in INR, inclusive of professional fees only (government "
    "fees stated separately where applicable)."
)

pricing = [
    ("6.1   Start a Business",
     ["Private Limited Company registration", "LLP registration", "OPC / Partnership / Proprietorship registration"],
     "one-time"),
    ("6.2   Statutory Registrations",
     ["GST registration", "MSME / Udyam registration", "Shops & Establishment / Trade license"],
     "one-time"),
    ("6.3   Licenses",
     ["FSSAI license", "Import Export Code (IEC)", "Other industry-specific license (specify)"],
     "one-time"),
    ("6.4   ROC Compliance",
     ["Annual ROC filing — Private Limited", "Annual ROC filing — LLP", "One-time event filings (specify)"],
     "yearly"),
    ("6.5   Tax Services",
     ["Income Tax Return — individual / proprietor", "Income Tax Return — company / LLP", "TDS return filing (quarterly)"],
     "yearly"),
    ("6.6   GST Services",
     ["GST registration (one-time)", "Monthly GST return filing (GSTR-1 and GSTR-3B)", "Annual GST return + GSTR-9C"],
     "as marked"),
    ("6.7   IPR (Intellectual Property)",
     ["Trademark filing — single class", "Copyright registration", "Patent filing (provisional)"],
     "one-time"),
    ("6.8   ISO Certification",
     ["ISO 9001 — Quality Management", "ISO 27001 — Information Security", "Annual surveillance audit"],
     "as marked"),
]

for service_label, tiers, frequency in pricing:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(service_label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    sub = p.add_run(f"   ({frequency} pricing)")
    sub.italic = True
    sub.font.size = Pt(9)
    sub.font.color.rgb = GREY

    t = add_field_table(doc)
    for j, tier in enumerate(tiers, 1):
        add_field_row(t, f"{service_label.split()[0]}.{j}   {tier}",
                      "Amount in INR. Example: Rs. 6,999. Or write \"On request\".")


# ===================================================================
# SECTION 7 — BLOG CONTENT
# ===================================================================
add_heading_block(doc, "7", "Blog content")
add_intro_para(doc,
    "The blog index page currently shows six placeholder articles. We have "
    "drafted five article topics that match what your customers search for. "
    "Confirm or replace each, and tell us which ones you want written first."
)
t = add_field_table(doc)
add_field_row(t, "7.1   Article 1 — \"GST registration documents and process\" — approve? (Y/N)",
              "If no, write your alternative topic.")
add_field_row(t, "7.2   Article 2 — \"ROC annual filing deadlines and penalties\" — approve? (Y/N)")
add_field_row(t, "7.3   Article 3 — \"Pvt Ltd vs LLP: which to choose\" — approve? (Y/N)")
add_field_row(t, "7.4   Article 4 — \"MSME Udyam registration process\" — approve? (Y/N)")
add_field_row(t, "7.5   Article 5 — \"ISO 9001 cost and timeline in India\" — approve? (Y/N)")
add_field_row(t, "7.6   Who will write the articles?",
              "Tick one: Our agency writes (additional charge) / Client provides drafts / Both / Decide later.")
add_field_row(t, "7.7   Target publish date for the first article",
              "Example: 15 July 2026.")


# ===================================================================
# SECTION 8 — VISUAL ASSETS
# ===================================================================
add_heading_block(doc, "8", "Visual assets to provide")
add_intro_para(doc,
    "The website currently uses generic stock images. To feel real and "
    "professional, we need photos and logos specific to your business. "
    "Tick each item the client agrees to provide and note the deadline."
)
add_checkbox_line(doc, "Hero photo for the homepage",
                  "professional, India-relevant — team at work, office shot, or composed scene")
add_checkbox_line(doc, "About page photo",
                  "founder portrait or team group photo")
add_checkbox_line(doc, "Photos of 4 team members",
                  "front-facing portraits, plain background")
add_checkbox_line(doc, "Office building or interior photo",
                  "for the contact page")
add_checkbox_line(doc, "Real client logos (6 needed for the trust carousel)",
                  "with each client's written permission to display")
add_checkbox_line(doc, "Case study photos (3 — one per case study above)",
                  "optional but improves credibility")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("8.X   Date by which photos will be shared with us: ")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY
r2 = p.add_run("_______ / _______ / 20_______")
r2.font.size = Pt(11)


# ===================================================================
# SECTION 9 — CONTACT FORM & EMAIL ROUTING
# ===================================================================
add_heading_block(doc, "9", "Contact form and email routing — DECISIONS")
add_intro_para(doc,
    "Right now the contact forms on all 10 service pages plus the contact "
    "page are designed but not yet connected to any inbox. Decide where "
    "form submissions should go and what should happen automatically."
)

add_intro_para(doc, "9.1   Where do form leads land first? (tick one)")
add_checkbox_line(doc, "A simple email to a fixed address",
                  "fastest setup — recommended if you do not yet have a CRM")
add_checkbox_line(doc, "Google Sheets (each lead becomes a new row)")
add_checkbox_line(doc, "A CRM (Zoho, HubSpot, Notion, etc.) — specify which: _________________________________")

add_intro_para(doc, "")
t = add_field_table(doc)
add_field_row(t, "9.2   Email address that should receive every new lead",
              "Example: leads@requiredfilings.com. Will receive an instant notification per form submission.")
add_field_row(t, "9.3   Auto-reply to the customer — what should it say?",
              "Default: \"Thanks for contacting RequiredFilings.com. We will reply within one working day.\" "
              "Edit or accept as-is.")
add_field_row(t, "9.4   Newsletter signups — where should those go?",
              "Same destination as form leads, or a separate list? If Mailchimp / Zoho Campaigns / Brevo etc., specify which.")


# ===================================================================
# SECTION 10 — OFFICE LOCATION (GOOGLE MAP)
# ===================================================================
add_heading_block(doc, "10", "Office location for Google Map")
add_intro_para(doc,
    "The contact page currently shows a placeholder map of Melbourne. "
    "Confirm the address Google should pin and whether you want it linked "
    "to a Google Business Profile."
)
t = add_field_table(doc)
add_field_row(t, "10.1  Exact address to pin on the map",
              "Same as 1.5 above, or different? Confirm.")
add_field_row(t, "10.2  Google Business Profile URL (if you have one)",
              "Example: https://g.page/your-business. We will link the map to it.")


# ===================================================================
# SECTION 11 — DOMAIN, HOSTING & SSL — DECISIONS
# ===================================================================
add_heading_block(doc, "11", "Domain, hosting and SSL — DECISIONS")
add_intro_para(doc,
    "The website needs to be hosted and pointed to your domain. These are "
    "decisions only the business owner can make."
)

add_intro_para(doc, "11.1  Domain registrar — where is requiredfilings.com registered? (tick one)")
add_checkbox_line(doc, "GoDaddy")
add_checkbox_line(doc, "Namecheap")
add_checkbox_line(doc, "BigRock")
add_checkbox_line(doc, "Google Domains / Squarespace")
add_checkbox_line(doc, "Other (specify): _________________________________")
add_checkbox_line(doc, "Not yet purchased — we want our agency to buy it on our behalf")

add_intro_para(doc, "")
t = add_field_table(doc)
add_field_row(t, "11.2  Login email for the domain registrar account",
              "Required so we can update DNS records to point the domain at the new website.")
add_field_row(t, "11.3  Authorised person to give DNS access to (name + email)",
              "We will coordinate the DNS change with this person. They must have admin access.")

add_intro_para(doc, "11.4  Hosting preference (tick one — first option recommended)")
add_checkbox_line(doc, "Our agency picks the best option for a static HTML site",
                  "recommended — typically Netlify, Vercel, or Cloudflare Pages. Free tier covers a brochure site.")
add_checkbox_line(doc, "Client already has hosting (specify provider): _________________________________")
add_checkbox_line(doc, "Client wants traditional cPanel hosting",
                  "more expensive, slower — only if there is a specific reason")

t = add_field_table(doc)
add_field_row(t, "11.5  Target go-live date",
              "Example: 30 July 2026. We need at least 5 working days after receiving this completed form.")


# ===================================================================
# SECTION 12 — ANALYTICS, TRACKING & COOKIES — DECISIONS
# ===================================================================
add_heading_block(doc, "12", "Analytics, tracking and cookies — DECISIONS")
add_intro_para(doc,
    "We strongly recommend installing analytics before launch so you can "
    "see what is working. All of the below are free."
)
add_checkbox_line(doc, "Install Google Analytics 4",
                  "recommended — shows visitor numbers, traffic sources, conversions")
add_checkbox_line(doc, "Install Google Search Console + submit sitemap",
                  "recommended — shows which keywords bring traffic")
add_checkbox_line(doc, "Install Microsoft Clarity (free heatmaps and session recordings)",
                  "recommended — see exactly how visitors use the site")
add_checkbox_line(doc, "Install Meta Pixel (Facebook / Instagram ad tracking)",
                  "tick only if you plan to run Meta ads")
add_checkbox_line(doc, "Install Google Ads conversion tracking",
                  "tick only if you plan to run Google ads")
add_checkbox_line(doc, "Show a cookie consent banner",
                  "required if any analytics or ad pixel above is installed")

add_intro_para(doc, "")
t = add_field_table(doc)
add_field_row(t, "12.X  Google account email for Analytics and Search Console",
              "If you do not have one, we can create a new one for you under your name.")


# ===================================================================
# SECTION 13 — ANYTHING ELSE
# ===================================================================
add_heading_block(doc, "13", "Anything else the client wants us to know")
add_intro_para(doc,
    "Special requests, brand sensitivities, competitors to avoid mentioning, "
    "preferred response tone, upcoming campaigns, or anything we have missed."
)
t = add_field_table(doc)
add_field_row(t, "13.1  Open notes")
add_field_row(t, " ")
add_field_row(t, " ")
add_field_row(t, " ")


# ===================================================================
# SIGN-OFF
# ===================================================================
add_heading_block(doc, "X", "Confirmation and sign-off")
add_intro_para(doc,
    "By signing below the client confirms that the information provided in "
    "this document is accurate and authorises RequiredFilings.com to "
    "publish it on the website, and to proceed with hosting, domain "
    "configuration and analytics setup as ticked above."
)

sig = doc.add_table(rows=2, cols=2)
sig.autofit = False
sig.columns[0].width = Cm(8)
sig.columns[1].width = Cm(8)

labels = [
    ("Signature of authorised signatory", "Date"),
    ("Name and designation (in print)", "Place"),
]
for ri, (l, r) in enumerate(labels):
    for ci, label in enumerate((l, r)):
        c = sig.rows[ri].cells[ci]
        c.text = ""
        p = c.paragraphs[0]
        rr = p.add_run("\n\n\n_______________________________________\n")
        rr.font.size = Pt(10)
        rrl = p.add_run(label)
        rrl.bold = True; rrl.font.size = Pt(9); rrl.font.color.rgb = GREY
        set_cell_border(c)

# footer note
fnp = doc.add_paragraph()
fnp.paragraph_format.space_before = Pt(18)
fnp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fnp.add_run("Once complete, please hand this document back to the team member who delivered it, "
                 "or scan and email it. The website goes live within 5 working days of receipt.")
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GREY


# ---------- save ----------
out = r"d:\Website Projects\Required Filings Website\client-intake-checklist.docx"
doc.save(out)
print(f"Saved: {out}")
